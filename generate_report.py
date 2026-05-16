import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_audit_report(filename):
    doc = docx.Document()

    # Título Principal
    title = doc.add_heading('Auditoría Técnica - Arquitectura Backend & DBA', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Proyecto: Ruta Fácil', style='Subtitle')
    doc.add_paragraph('Módulo Auditado: backend/logistics\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sección 1: Análisis de Modelos de Base de Datos (models.py)
    doc.add_heading('1. Arquitectura de Datos (models.py)', level=1)
    doc.add_paragraph(
        "Se revisó el esquema de base de datos actual. A continuación, los hallazgos y áreas de mejora:"
    )
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Modelos Existentes: ").bold = True
    p.add_run("Se encuentran definidos Aliado, Repartidor, Cliente, Pedido y Ruta, cumpliendo con la estructura básica solicitada.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Seguridad y RBAC: ").bold = True
    p.add_run("Actualmente, Aliado y Repartidor tienen una relación OneToOne con el modelo User nativo de Django, pero falta un control explícito de roles. Se recomienda crear una extensión de User o utilizar Groups nativos para definir explícitamente los roles 'admin', 'aliado' y 'repartidor' con permisos DRF.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Geolocalización: ").bold = True
    p.add_run("Las coordenadas latitud y longitud están correctamente tipadas como DecimalField, lo cual es ideal para precisión espacial.")

    # Sección 2: Vistas y Algoritmia (views.py)
    doc.add_heading('2. Lógica de Negocio y Algoritmia (views.py)', level=1)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Algoritmo de Asignación: ").bold = True
    p.add_run("Se implementó un algoritmo Greedy de emparejamiento (Vecino más cercano) en PedidoViewSet.asignar_automatico. La complejidad actual es de O(M*N), lo cual es aceptable para volumenes medios. Utiliza distancia euclidiana, que puede ser mejorada a distancia Haversine si se requiere mayor precisión geográfica.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Protección de Endpoints (Security by Design): ").bold = True
    run = p.add_run(" CRÍTICO. ")
    run.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run("Actualmente las vistas (ViewSets) carecen de los decoradores o clases de permisos (permission_classes = [IsAuthenticated, ...]). Cualquier usuario podría invocar 'asignar_automatico' o modificar datos.")

    # Sección 3: Serializadores y API (serializers.py)
    doc.add_heading('3. Integración con Frontend (serializers.py)', level=1)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Overfetching/Underfetching: ").bold = True
    p.add_run("El PedidoSerializer utiliza ReadOnlyFields correctamente para anidar nombres y coordenadas del cliente y el repartidor sin necesidad de enviar peticiones adicionales, lo que es una buena práctica de integración con React.")

    # Sección 4: Plan de Acción y Correcciones Propuestas
    doc.add_heading('4. Plan de Acción y Siguientes Pasos', level=1)
    
    doc.add_paragraph("Para cumplir con los lineamientos del rol de Arquitecto Backend, se recomienda ejecutar las siguientes modificaciones:")
    
    doc.add_paragraph("1. Implementar Type Hints (tipado fuerte) en los métodos de models.py y views.py.", style='List Number')
    doc.add_paragraph("2. Configurar permissions y Custom Permissions en DRF para garantizar el Control de Acceso Basado en Roles (RBAC).", style='List Number')
    doc.add_paragraph("3. Refinar el algoritmo O(M*N) implementando el cálculo Haversine para distancias terrestres y envolviendo el proceso en una transacción atómica (@transaction.atomic) para evitar asignaciones duplicadas.", style='List Number')

    doc.save(filename)
    print(f"Reporte generado exitosamente en: {filename}")

if __name__ == '__main__':
    create_audit_report('Auditoria_Backend_Ruta_Facil.docx')
