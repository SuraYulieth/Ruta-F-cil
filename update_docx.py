import docx

def update_document(input_filename, output_filename):
    doc = docx.Document(input_filename)
    
    # Textos a inyectar
    algoritmo_texto = " Algoritmo Ávido o Greedy Matching (Vecino Más Cercano) para la asignación de rutas y repartidores."
    lenguaje_texto = " Python 3 con el framework Django REST Framework (DRF) y Base de datos SQLite3."
    mejor_caso_texto = " O(M) - Ocurre cuando el primer repartidor evaluado está disponible y cumple las condiciones, iterando una sola vez por pedido."
    caso_promedio_texto = " O(M x N) - Se evalúan múltiples repartidores libres para cada pedido en tránsito."
    peor_caso_texto = " O(M x N) - El algoritmo recorre todos los repartidores para cada pedido sin encontrar coincidencias rápidas."
    
    codigo_texto = """@action(detail=False, methods=['post'])
def asignar_automatico(self, request):
    pedidos_pendientes = Pedido.objects.filter(estado='Pendiente')
    repartidores_libres = CustomUser.objects.filter(role='driver', estado='Disponible')

    asignaciones = 0
    for pedido in pedidos_pendientes:
        mejor_repartidor = repartidores_libres.first() # Greedy

        if mejor_repartidor:
            pedido.repartidor = mejor_repartidor
            pedido.estado = 'Asignado'
            pedido.save()
            mejor_repartidor.estado = 'Ocupado'
            mejor_repartidor.save()
            asignaciones += 1

    return Response({"mensaje": f"Se asignaron {asignaciones} pedidos."})"""

    # Buscar y reemplazar
    for p in doc.paragraphs:
        if "Algoritmo utilizado:" in p.text:
            p.add_run(algoritmo_texto)
        elif "Lenguaje de programación:" in p.text:
            p.add_run(lenguaje_texto)
        elif "Mejor caso:" in p.text:
            p.add_run(mejor_caso_texto)
        elif "Caso promedio:" in p.text:
            p.add_run(caso_promedio_texto)
        elif "Peor caso:" in p.text:
            p.add_run(peor_caso_texto)
        elif "Ejemplo de código desarrollado:" in p.text:
            p.insert_paragraph_before("").add_run(codigo_texto)

    doc.save(output_filename)
    print(f"Documento guardado como: {output_filename}")

if __name__ == '__main__':
    update_document('Ruta Fácil _Proyecto_ABP (1) (1).docx', 'Ruta Fácil _Proyecto_ABP_Completado.docx')
