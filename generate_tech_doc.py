import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_tech_doc(filename):
    doc = docx.Document()

    # Título Principal
    title = doc.add_heading('Documentación Técnica - Backend Ruta Fácil', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sección 1
    doc.add_heading('1. Información General', level=1)
    doc.add_paragraph('El backend del proyecto "Ruta Fácil" está desarrollado bajo una arquitectura RESTful robusta orientada a proveer servicios al frontend en React, delegando toda la complejidad transaccional y de asignación (lógica de negocio) al servidor.')
    
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run('Lenguaje: ').bold = True
    p1.add_run('Python 3')
    
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run('Framework Principal: ').bold = True
    p2.add_run('Django 6.0')
    
    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run('Framework API: ').bold = True
    p3.add_run('Django REST Framework (DRF)')
    
    p4 = doc.add_paragraph(style='List Bullet')
    p4.add_run('Base de Datos: ').bold = True
    p4.add_run('SQLite (db.sqlite3)')
    
    p5 = doc.add_paragraph(style='List Bullet')
    p5.add_run('Módulo Principal: ').bold = True
    p5.add_run('logistics')

    # Sección 2
    doc.add_heading('2. Configuración e Instalación', level=1)
    doc.add_paragraph('Para ejecutar este entorno en local, se deben seguir los siguientes pasos:')
    
    doc.add_paragraph('1. Activar Entorno Virtual:', style='List Number')
    doc.add_paragraph('.\\venv\\Scripts\\activate', style='Intense Quote')
    
    doc.add_paragraph('2. Instalar Dependencias:', style='List Number')
    doc.add_paragraph('pip install django djangorestframework django-cors-headers', style='Intense Quote')
    
    doc.add_paragraph('3. Ejecutar Migraciones:', style='List Number')
    doc.add_paragraph('python manage.py makemigrations\npython manage.py migrate', style='Intense Quote')
    
    doc.add_paragraph('4. Sembrar Datos Iniciales (Mock Data):', style='List Number')
    doc.add_paragraph('python seed.py', style='Intense Quote')
    
    doc.add_paragraph('5. Levantar el Servidor:', style='List Number')
    doc.add_paragraph('python manage.py runserver', style='Intense Quote')

    # Sección 3
    doc.add_heading('3. Arquitectura de Datos (Modelos)', level=1)
    doc.add_paragraph('El motor principal reside en el modelo CustomUser y su relación estricta con el ecosistema de distribución.')
    
    p_usr = doc.add_paragraph(style='List Bullet')
    p_usr.add_run('CustomUser (Control de Acceso - RBAC): ').bold = True
    p_usr.add_run('Extiende el usuario nativo de Django. Centraliza credenciales y roles. Campos clave: role, nombre, estado, ubicacion.')
    
    p_cli = doc.add_paragraph(style='List Bullet')
    p_cli.add_run('Cliente: ').bold = True
    p_cli.add_run('Almacena la información de los usuarios finales (compradores). Campos clave: nombre, direccion, latitud, longitud.')
    
    p_ped = doc.add_paragraph(style='List Bullet')
    p_ped.add_run('Pedido: ').bold = True
    p_ped.add_run('El núcleo transaccional. Vincula a un Cliente (obligatorio) y dinámicamente a un CustomUser (repartidor).')

    # Sección 4
    doc.add_heading('4. Endpoints de la API (Rutas RESTful)', level=1)
    
    doc.add_heading('4.1 Autenticación', level=2)
    doc.add_paragraph('Ruta: POST /api/login/', style='List Bullet')
    doc.add_paragraph('Descripción: Valida credenciales y devuelve el perfil completo.', style='List Bullet')
    
    doc.add_heading('4.2 Gestión de Usuarios', level=2)
    doc.add_paragraph('Ruta: GET /api/users/', style='List Bullet')
    doc.add_paragraph('Descripción: Lista todos los usuarios registrados (repartidores y administradores), aplanando el modelo para que concuerde con React.', style='List Bullet')
    
    doc.add_heading('4.3 Gestión de Pedidos', level=2)
    doc.add_paragraph('Ruta: GET /api/pedidos/ | POST /api/pedidos/ | PATCH /api/pedidos/{id}/', style='List Bullet')
    doc.add_paragraph('Descripción: Permite al frontend listar, crear y actualizar un pedido.', style='List Bullet')
    
    doc.add_heading('4.4 Algoritmia de Asignación', level=2)
    doc.add_paragraph('Ruta: POST /api/pedidos/asignar_automatico/', style='List Bullet')
    doc.add_paragraph('Descripción: Algoritmo de emparejamiento Greedy O(M x N).', style='List Bullet')

    # Sección 5
    doc.add_heading('5. Decisiones de Diseño y Seguridad', level=1)
    doc.add_paragraph('Prevención de Overfetching: Se sobrescribieron los métodos nativos de DRF para formatear la respuesta HTTP en la misma estructura estricta del context de React.', style='List Bullet')
    doc.add_paragraph('RBAC Seguro: Se inyectó una entidad global con roles validados desde la Base de Datos.', style='List Bullet')

    # Guardar
    doc.save(filename)
    print(f"Reporte generado exitosamente en: {filename}")

if __name__ == '__main__':
    os.makedirs('documentacion', exist_ok=True)
    create_tech_doc('documentacion/Documentacion_Tecnica_Backend.docx')
